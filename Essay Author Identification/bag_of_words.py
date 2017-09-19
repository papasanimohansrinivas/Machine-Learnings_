authors = {}

docs =0

tril = 0

lines = 0

Package_not_found =0

KeyError_ = 0

UnicodeEncodeError_ = 0 

persons_lines = {}


def get_roll_number(roll_numbers,files):

	id_ = None
	for number in roll_numbers:
		if number in files:
			id_ = number
			return id_
	return id_


def text_extraction():
	import textract

	import script 

	import os

	import docx

	import zipfile
	import pyth
	
	from pyth.plugins.rtf15.reader import Rtf15Reader

	from pyth.plugins.plaintext.writer import PlaintextWriter


	roll_numbers = script.roll_numbers

	file_names = script.file_names


	# import subprocess

	Document = docx.Document

	global docs

	global authors
	global tril 
	global Package_not_found
	global KeyError_
	global UnicodeEncodeError_

	global lines

	one = 0 
	two = 0
	three = 0

	for folders in file_names:
		for files in os.listdir(folders):	
			docs+=1
			file_format = files[len(files)-4:]
			file_name = os.getcwd()+"\ ".strip()+folders+"\ ".strip()+files

			try:
				id_= get_roll_number(roll_numbers,files)

				text_ = None

				if file_format==".odt":
					text_ = textract.process(file_name,extension='odt')
					one+=1

				elif file_format=="docx":
					document = Document(file_name)

					text_ = ' '.join(paragraph.text.encode('utf-8') for paragraph in document.paragraphs)
					two+=1
				elif file_format==".zip":
					fi = zipfile.ZipFile(file_name,"r")
					three+=1
					for names in fi.namelist():
						if names[len(names)-4:]=="docx":
							doc=Document(names)
							text_ = ' '.join(par.text.encode('utf-8') for par in doc.paragraphs)
							print text_
				elif file_format==".rar":
					# zipfile.ZipFile(file_name,"r")
					# print "rar"
					pass
				elif file_format==".txt":
					file_ = open(file_name,"r")
					text_ = file_.read()
					# print type(text_) 
				elif file_format==".rtf":
					# rtf_file = open(file_name,"rb")
					try:
						doc   =  Rtf15Reader.read(open(file_name))
						text_ =  PlaintextWriter.write(doc).getvalue()
					# print folders,files
					# textract.process(file_name)
					# print "--------------"
					# print text_
					except pyth.errors.WrongFileType:
						print files,folders
				else:
					lines+=1
					# print file_format

			except docx.opc.exceptions.PackageNotFoundError:
				# print "Package NOt Found Error",files[len(files)-3:]
				Package_not_found+=1

	print one,two,three,[one+two+three]

 
text_extraction()


import PyRTF

# help(PyRTF)
# print KeyError_,Package_not_found,UnicodeEncodeError_,sum([KeyError_,Package_not_found,UnicodeEncodeError_])
print lines,Package_not_found

# print docs

# print tril

ser = 0

length_text = 0

def write_text_to_file():
	global length_text
	global ser
	for name in authors:
		if len(authors[name])==0:
			ser+=1
		length_text+=len(authors[name].split())
		if name!=None:
			z=open(name+".txt","w")
			z.write(authors[name])
			z.close()

# write_text_to_file()